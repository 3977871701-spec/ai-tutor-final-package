import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from pypdf import PdfReader
from docx import Document


def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    docs = []
    reader = PdfReader(file_path)
    full_text = ""
    
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            full_text += f"\n--- 第{i+1}页 ---\n{text}"
    
    # Split by sections/paragraphs for better chunking
    chunks = split_into_chunks(full_text)
    
    for chunk in chunks:
        if len(chunk.strip()) > 20:  # Skip very short chunks
            docs.append({
                "content": chunk.strip(),
                "metadata": {
                    "source": file_path,
                    "type": "pdf",
                    "file_name": Path(file_path).name
                }
            })
    
    return docs


def parse_word(file_path: str) -> List[Dict[str, Any]]:
    docs = []
    doc = Document(file_path)
    full_text = ""
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text.strip() + "\n"
    
    # Process tables
    for i, table in enumerate(doc.tables):
        table_text = f"\n--- 表格 {i+1} ---\n"
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            table_text += row_text + "\n"
        full_text += table_text
    
    chunks = split_into_chunks(full_text)
    
    for chunk in chunks:
        if len(chunk.strip()) > 20:
            docs.append({
                "content": chunk.strip(),
                "metadata": {
                    "source": file_path,
                    "type": "word",
                    "file_name": Path(file_path).name
                }
            })
    
    return docs


def parse_text(content: str, metadata: Optional[Dict] = None) -> List[Dict[str, Any]]:
    chunks = split_into_chunks(content)
    meta = metadata or {}
    return [{
        "content": chunk.strip(),
        "metadata": {**meta, "type": "text"}
    } for chunk in chunks if len(chunk.strip()) > 20]


def split_into_chunks(text: str, max_chunk_size: int = 500, overlap: int = 50) -> List[str]:
    # Split by double newlines (paragraphs)
    paragraphs = re.split(r'\n\s*\n', text)
    
    chunks = []
    current_chunk = ""
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        if len(current_chunk) + len(para) < max_chunk_size:
            current_chunk += para + "\n\n"
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # Start new chunk with overlap
            words = current_chunk.split()
            if len(words) > overlap // 5 and overlap > 0:
                # Keep some overlap
                overlap_words = words[-(overlap // 5):]
                current_chunk = " ".join(overlap_words) + "\n\n" + para + "\n\n"
            else:
                current_chunk = para + "\n\n"
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks


def parse_document(file_path: str) -> List[Dict[str, Any]]:
    suffix = Path(file_path).suffix.lower()
    
    if suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return parse_word(file_path)
    elif suffix == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return parse_text(f.read(), {"source": file_path, "file_name": Path(file_path).name})
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
